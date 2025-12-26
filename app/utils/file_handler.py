"""
File handling utilities
Hỗ trợ: TXT, PDF, MD, DOCX, DOC
"""
from pathlib import Path
from typing import List
import logging
import docx
import docx2txt
from langchain.schema import Document
from langchain_community.document_loaders import TextLoader, PyPDFLoader, UnstructuredMarkdownLoader

logger = logging.getLogger(__name__)


class FileHandler:
    """Utility class for handling different file types"""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.md', '.docx', '.doc'}
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file extension is supported"""
        ext = Path(filename).suffix.lower()
        return ext in FileHandler.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """Get file type from filename"""
        return Path(filename).suffix.lower().replace('.', '')
    
    @staticmethod
    def load_document(file_path: Path) -> List[Document]:
        """
        Load document based on file extension
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of Document objects
            
        Raises:
            ValueError: If file type is not supported
        """
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.txt':
                return FileHandler._load_txt(file_path)
            elif extension == '.pdf':
                return FileHandler._load_pdf(file_path)
            elif extension == '.md':
                return FileHandler._load_markdown(file_path)
            elif extension == '.docx':
                return FileHandler._load_docx(file_path)
            elif extension == '.doc':
                return FileHandler._load_doc(file_path)
            else:
                raise ValueError(f"Định dạng file không được hỗ trợ: {extension}")
                
        except Exception as e:
            logger.error(f"Lỗi khi load document {file_path.name}: {str(e)}")
            raise
    
    @staticmethod
    def _load_txt(file_path: Path) -> List[Document]:
        """Load text file"""
        loader = TextLoader(str(file_path), encoding='utf-8')
        documents = loader.load()
        logger.info(f"Đã load {len(documents)} pages từ {file_path.name}")
        return documents
    
    @staticmethod
    def _load_pdf(file_path: Path) -> List[Document]:
        """Load PDF file"""
        loader = PyPDFLoader(str(file_path))
        documents = loader.load()
        logger.info(f"✅ Đã load {len(documents)} pages từ {file_path.name}")
        return documents
    
    @staticmethod
    def _load_markdown(file_path: Path) -> List[Document]:
        """Load Markdown file"""
        loader = UnstructuredMarkdownLoader(str(file_path))
        documents = loader.load()
        logger.info(f"✅ Đã load {len(documents)} pages từ {file_path.name}")
        return documents
    
    @staticmethod
    def _load_docx(file_path: Path) -> List[Document]:
        """Load DOCX file"""
        try:
            # Try using python-docx first
            doc = docx.Document(str(file_path))
            text = '\n\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            if not text.strip():
                # Fallback to docx2txt
                text = docx2txt.process(str(file_path))
            
            document = Document(
                page_content=text,
                metadata={"source": str(file_path), "filename": file_path.name}
            )
            logger.info(f"Đã load DOCX file: {file_path.name}")
            return [document]
            
        except Exception as e:
            logger.error(f"Lỗi khi load DOCX: {str(e)}")
            raise
    
    @staticmethod
    def _load_doc(file_path: Path) -> List[Document]:
        """Load DOC file (old Word format)"""
        try:
            # Use docx2txt for .doc files
            text = docx2txt.process(str(file_path))
            
            if not text.strip():
                raise ValueError("Không thể extract text từ file .doc")
            
            document = Document(
                page_content=text,
                metadata={"source": str(file_path), "filename": file_path.name}
            )
            logger.info(f"Đã load DOC file: {file_path.name}")
            return [document]
            
        except Exception as e:
            logger.error(f"Lỗi khi load DOC: {str(e)}")
            raise
    
    @staticmethod
    def validate_file_size(file_size: int, max_size: int) -> bool:
        """Validate file size"""
        return file_size <= max_size
    
    @staticmethod
    def format_file_size(size_bytes: int) -> str:
        """Format file size to human readable string"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"
